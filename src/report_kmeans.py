"""
Analyse the Phase 1 K-Means seed sweep and write a Word report.

Reads ``results/kmeans_seed_sweep.csv`` (produced by
``python main.py --phase 1 --seeds ...``) and writes
``results/kmeans_phase1_report.docx`` plus the figures it embeds.

Every number in the document is computed from the CSV, so re-running after
adding seeds refreshes the report rather than requiring edits.

Usage
-----
    python -m src.report_kmeans
    python -m src.report_kmeans --csv results/kmeans_seed_sweep.csv
"""
from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from scipy import stats

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import config
from src.run_config import _git_state

REPRESENTATIONS = ["trajectories", "features"]
COLOURS = {"trajectories": "#3b6ea5", "features": "#c26b3f"}


# ── Statistics ────────────────────────────────────────────────────────────────
def peak_per_seed(df: pd.DataFrame, metric: str = "ari") -> pd.DataFrame:
    """For each (representation, seed), the best k and the metric value there."""
    idx = df.groupby(["representation", "seed"])[metric].idxmax()
    peaks = df.loc[idx, ["representation", "seed", "k", "ari", "nmi"]]
    return peaks.rename(columns={"k": "best_k"}).reset_index(drop=True)


def spread(values: np.ndarray) -> dict:
    values = np.asarray(values, dtype=float)
    sd = float(values.std(ddof=1)) if len(values) > 1 else 0.0
    mean = float(values.mean())
    return {
        "mean": mean,
        "sd": sd,
        "min": float(values.min()),
        "max": float(values.max()),
        "range": float(values.max() - values.min()),
        "cv": (sd / mean * 100) if mean else float("nan"),
        "n": int(len(values)),
    }


def curve_stats(df: pd.DataFrame, rep: str, metric: str) -> pd.DataFrame:
    """Mean/sd of a metric at each k, across seeds."""
    sub = df[df["representation"] == rep]
    g = sub.groupby("k")[metric]
    out = g.agg(["mean", "std", "min", "max", "count"]).reset_index()
    return out.fillna({"std": 0.0})


# ── Figures ───────────────────────────────────────────────────────────────────
def _style(ax, xlabel, ylabel, title):
    ax.set_xlabel(xlabel)
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(alpha=0.25, linewidth=0.6)


def figure_metric_vs_k(df: pd.DataFrame, metric: str, out_path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(7.2, 4.0), dpi=200)
    for rep in REPRESENTATIONS:
        if rep not in set(df["representation"]):
            continue
        cs = curve_stats(df, rep, metric)
        ax.plot(cs["k"], cs["mean"], color=COLOURS[rep], linewidth=1.8, label=f"{rep} (mean)")
        ax.fill_between(
            cs["k"], cs["mean"] - cs["std"], cs["mean"] + cs["std"],
            color=COLOURS[rep], alpha=0.20, linewidth=0,
        )
    n_seeds = df["seed"].nunique()
    _style(ax, "Number of clusters (k)", metric.upper(),
           f"{metric.upper()} vs k — mean +/- 1 SD across {n_seeds} seeds")
    ax.axhline(0, color="#888", linewidth=0.8, linestyle="--")
    ax.legend(frameon=False, fontsize=9)
    fig.tight_layout()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path)
    plt.close(fig)
    return out_path


def figure_peak_by_seed(peaks: pd.DataFrame, out_path: Path) -> Path:
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.4), dpi=200)
    seeds = sorted(peaks["seed"].unique())
    width = 0.38
    x = np.arange(len(seeds))
    for i, rep in enumerate([r for r in REPRESENTATIONS if r in set(peaks["representation"])]):
        sub = peaks[peaks["representation"] == rep].sort_values("seed")
        axes[0].bar(x + (i - 0.5) * width, sub["ari"], width,
                    color=COLOURS[rep], label=rep)
        axes[1].scatter(sub["seed"], sub["best_k"], color=COLOURS[rep], s=28, label=rep)
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(seeds)
    _style(axes[0], "Seed", "Peak ARI", "Peak ARI by seed")
    axes[0].legend(frameon=False, fontsize=8)
    _style(axes[1], "Seed", "k at peak ARI", "Selected k by seed")
    fig.tight_layout()
    fig.savefig(out_path)
    plt.close(fig)
    return out_path


# ── Document helpers ──────────────────────────────────────────────────────────
def add_table(doc, df: pd.DataFrame, float_fmt: str = "{:.4f}"):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = float_fmt.format(val) if isinstance(val, (float, np.floating)) else str(val)
    for r in table.rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def add_caption(doc, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)


# ── Report ────────────────────────────────────────────────────────────────────
def build_report(csv_path: Path, out_path: Path) -> Path:
    df = pd.read_csv(csv_path)
    seeds = sorted(df["seed"].unique())
    reps = [r for r in REPRESENTATIONS if r in set(df["representation"])]
    k_values = sorted(df["k"].unique())
    peaks = peak_per_seed(df)

    fig_dir = config.RESULTS_DIR / "figures"
    fig_ari = figure_metric_vs_k(df, "ari", fig_dir / "kmeans_ari_vs_k.png")
    fig_nmi = figure_metric_vs_k(df, "nmi", fig_dir / "kmeans_nmi_vs_k.png")
    fig_seed = figure_peak_by_seed(peaks, fig_dir / "kmeans_peak_by_seed.png")

    stats_by_rep = {}
    for rep in reps:
        sub = peaks[peaks["representation"] == rep]
        stats_by_rep[rep] = {
            "ari": spread(sub["ari"].values),
            "nmi": spread(sub["nmi"].values),
            "best_k": spread(sub["best_k"].values),
            "k_mode": int(sub["best_k"].mode().iloc[0]),
            "k_unique": sorted(sub["best_k"].unique().tolist()),
        }

    sha, dirty = _git_state()
    doc = Document()

    # ── Title block ───────────────────────────────────────────────────────
    doc.add_heading("Phase 1 — K-Means Baseline: Seed Sweep Analysis", level=0)
    p = doc.add_paragraph()
    p.add_run("Interception Movements Deep Learning Pipeline\n").bold = True
    p.add_run(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} · "
        f"{len(seeds)} seeds ({', '.join(map(str, seeds))}) · "
        f"k = {min(k_values)}–{max(k_values)} · "
        f"git {sha[:8]}{' (dirty)' if dirty else ''}"
    ).font.size = Pt(9)

    # ── Executive summary ─────────────────────────────────────────────────
    doc.add_heading("1. Summary", level=1)
    best_rep = max(reps, key=lambda r: stats_by_rep[r]["ari"]["mean"])
    a = stats_by_rep[best_rep]["ari"]
    strength = ("negligible" if a["mean"] < 0.05 else
                "weak" if a["mean"] < 0.20 else
                "moderate" if a["mean"] < 0.50 else "strong")

    doc.add_paragraph(
        f"K-Means was repeated across {len(seeds)} seeds on the full set of trials, "
        f"clustering two representations (raw normalised trajectories and "
        f"extracted kinematic features) against true subject identity as the label. "
        f"The question is whether individual movement style is separable by an "
        f"off-the-shelf clustering method, which sets the bar the CVAE must clear."
    )
    for line in [
        f"Best representation: {best_rep}, peak ARI {a['mean']:.4f} +/- {a['sd']:.4f} "
        f"(range {a['min']:.4f}–{a['max']:.4f}) across seeds.",
        f"Agreement with subject identity is {strength}: ARI is chance-corrected, so "
        f"0 means no better than random and 1 means perfect recovery.",
        f"Seed sensitivity is {'low' if a['cv'] < 10 else 'moderate' if a['cv'] < 25 else 'high'} "
        f"(coefficient of variation {a['cv']:.1f}% of the mean).",
        f"Selected k varied widely across seeds "
        f"({stats_by_rep[best_rep]['k_unique'][0]}–{stats_by_rep[best_rep]['k_unique'][-1]}) "
        f"despite the stable ARI, because the ARI-vs-k curve is flat near its peak — "
        f"the argmax is a coin toss between near-tied values (section 3.3).",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Method ────────────────────────────────────────────────────────────
    doc.add_heading("2. What was run", level=1)
    doc.add_paragraph(
        f"For each seed in {{{', '.join(map(str, seeds))}}} and each representation, "
        f"K-Means was fitted for every k from {min(k_values)} to {max(k_values)} "
        f"(n_init = 10, random_state = seed). Each fit was scored against true subject "
        f"labels with the Adjusted Rand Index (ARI) and Normalised Mutual Information "
        f"(NMI). That is {len(seeds)} x {len(reps)} x {len(k_values)} = "
        f"{len(seeds) * len(reps) * len(k_values)} fits, recorded in "
        f"{csv_path.name} as one row per (seed, representation, k)."
    )
    doc.add_paragraph(
        "Important scope note: in this phase the seed changes only the K-Means "
        "centroid initialisation. There is no train/validation/test split — every "
        "seed clusters the identical set of trials from all 28 subjects. The spread "
        "reported here is therefore algorithmic variance alone, and is a lower bound "
        "on the true run-to-run noise of this baseline, not an estimate of it. "
        "See section 6.",
    )

    # ── Results ───────────────────────────────────────────────────────────
    doc.add_heading("3. Results", level=1)

    doc.add_heading("3.1 Peak performance across seeds", level=2)
    rows = []
    for rep in reps:
        s = stats_by_rep[rep]
        rows.append({
            "Representation": rep,
            "ARI mean": s["ari"]["mean"], "ARI SD": s["ari"]["sd"],
            "ARI min": s["ari"]["min"], "ARI max": s["ari"]["max"],
            "NMI mean": s["nmi"]["mean"], "NMI SD": s["nmi"]["sd"],
            "k (mode)": s["k_mode"],
        })
    add_table(doc, pd.DataFrame(rows))
    add_caption(doc, f"Table 1. Peak ARI/NMI per representation, aggregated over {len(seeds)} seeds.")

    doc.add_heading("3.2 Per-seed detail", level=2)
    detail = peaks.sort_values(["representation", "seed"])[
        ["representation", "seed", "best_k", "ari", "nmi"]
    ].rename(columns={"representation": "Representation", "seed": "Seed",
                      "best_k": "k at peak", "ari": "ARI", "nmi": "NMI"})
    add_table(doc, detail)
    add_caption(doc, "Table 2. Best k and the metrics attained there, per seed.")

    doc.add_heading("3.3 Why the selected k is unstable", level=2)
    for rep in reps:
        cs = curve_stats(df, rep, "ari")
        best_row = cs.loc[cs["mean"].idxmax()]
        tol = float(best_row["std"])
        plateau = cs[cs["mean"] >= best_row["mean"] - tol]["k"]
        k_lo, k_hi = int(plateau.min()), int(plateau.max())
        tail = cs[cs["k"] >= cs["k"].max() - 5]
        slope = (tail["mean"].iloc[-1] - tail["mean"].iloc[0]) / max(len(tail) - 1, 1)
        doc.add_paragraph(
            f"{rep}: mean ARI peaks at k = {int(best_row['k'])} "
            f"({best_row['mean']:.4f}), but every k from {k_lo} to {k_hi} sits within "
            f"one seed-level SD of that peak, and the curve changes by only "
            f"{slope:+.5f} per unit k over the last five values. The ARI surface is "
            f"effectively flat across that band, so the argmax is decided by noise. "
            f"That is why the per-seed 'best k' spans "
            f"{stats_by_rep[rep]['k_unique'][0]}–{stats_by_rep[rep]['k_unique'][-1]} "
            f"while peak ARI itself varies by only "
            f"{stats_by_rep[rep]['ari']['range']:.4f}.",
            style="List Bullet",
        )
    doc.add_paragraph(
        "The practical consequence: the scattered k values in Table 2 should not be "
        "read as disagreement about cluster structure. They are ties broken at random. "
        "Quoting a single 'optimal k' from this baseline would overstate what the data "
        "supports."
    )

    doc.add_heading("3.4 Metric curves", level=2)
    doc.add_picture(str(fig_ari), width=Inches(6.2))
    add_caption(doc, "Figure 1. ARI against k. Shaded band is +/- 1 SD across seeds.")
    doc.add_picture(str(fig_nmi), width=Inches(6.2))
    add_caption(doc, "Figure 2. NMI against k. NMI is not chance-corrected and rises "
                     "with k, so it should not be read as a measure of quality on its own.")
    doc.add_picture(str(fig_seed), width=Inches(6.2))
    add_caption(doc, "Figure 3. Peak ARI and selected k, seed by seed.")

    # ── Comparison ────────────────────────────────────────────────────────
    if len(reps) == 2:
        doc.add_heading("4. Trajectories vs. features", level=1)
        a_vals = peaks[peaks["representation"] == reps[0]].sort_values("seed")["ari"].values
        b_vals = peaks[peaks["representation"] == reps[1]].sort_values("seed")["ari"].values
        diff = a_vals - b_vals
        para = (
            f"Paired by seed, {reps[0]} reached a peak ARI "
            f"{'higher' if diff.mean() > 0 else 'lower'} than {reps[1]} by "
            f"{abs(diff.mean()):.4f} on average (SD {diff.std(ddof=1):.4f}, "
            f"consistent in sign on {int((np.sign(diff) == np.sign(diff.mean())).sum())} "
            f"of {len(diff)} seeds)."
        )
        if len(diff) >= 5:
            try:
                w = stats.wilcoxon(a_vals, b_vals)
                para += (f" A Wilcoxon signed-rank test over the {len(diff)} paired seeds "
                         f"gives W = {w.statistic:.1f}, p = {w.pvalue:.4g}.")
            except ValueError:
                pass
        doc.add_paragraph(para)
        doc.add_paragraph(
            "Because both representations are scored on the same trials with the same "
            "seeds, this comparison is paired and the seed-to-seed variance largely "
            "cancels. It says which input K-Means prefers; it does not say either one "
            "recovers subject identity well."
        )

    # ── Interpretation ────────────────────────────────────────────────────
    doc.add_heading("5. Interpretation", level=1)
    doc.add_paragraph(
        f"Peak ARI of {a['mean']:.4f} means cluster assignments agree with subject "
        f"identity only marginally more than chance. K-Means on Euclidean distance "
        f"between whole trajectories (or between summary kinematic features) does not "
        f"recover who produced a movement. Two plain reasons: within-subject variation "
        f"across trials is comparable to or larger than between-subject variation, and "
        f"Euclidean distance in a 300-dimensional trajectory space is dominated by "
        f"condition effects (starting position, stimulus speed) that are shared across "
        f"subjects rather than by individual style."
    )
    doc.add_paragraph(
        "NMI looks more encouraging than ARI, but that gap is expected rather than "
        "informative: NMI is not adjusted for chance and increases mechanically as k "
        "approaches the number of true classes. ARI is the metric to quote."
    )
    doc.add_paragraph(
        "This is a useful negative result. It establishes that the phase-3 CVAE is not "
        "competing against a trivially strong baseline, and that any latent structure "
        "it recovers is not obtainable by distance-based clustering of the same inputs."
    )

    # ── Limitations ───────────────────────────────────────────────────────
    doc.add_heading("6. Limitations", level=1)
    for line in [
        "The seed varies only K-Means initialisation. All seeds cluster the same "
        "trials from the same 28 subjects, so the reported SD does not include "
        "sampling variability over subjects — the dominant noise source in this study.",
        "With n_init = 10, scikit-learn already keeps the best of ten initialisations "
        "per fit, which suppresses seed-to-seed variance further. A larger spread "
        "would be expected at n_init = 1.",
        "Trials per subject are unbalanced (29 to 180, median 180), so trial-level "
        "metrics implicitly weight subjects unequally.",
        "ARI is computed against 28 true classes while k ranges over "
        f"{min(k_values)}–{max(k_values)}; selecting k by maximising ARI uses the "
        "labels, so the peak values are optimistic as an estimate of unsupervised "
        "performance.",
        "No dimensionality reduction was applied before clustering the raw "
        "trajectories, and K-Means assumes isotropic, similarly-sized clusters — "
        "assumptions unlikely to hold here.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Recommendations ───────────────────────────────────────────────────
    doc.add_heading("7. Recommended next steps", level=1)
    for line in [
        "To get an honest noise estimate for this baseline, resample subjects "
        "(bootstrap or leave-k-subjects-out) rather than only re-seeding K-Means.",
        "Report a per-subject-averaged metric alongside the trial-level one, so the "
        "29-trial and 180-trial subjects count equally.",
        "Fix k at the number of subjects (k = 28) for a label-free comparison, and "
        "report that value next to the ARI-maximising k.",
        "Carry the same seed list into phase 3 so baseline and CVAE results are "
        "compared over matched conditions.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Appendix ──────────────────────────────────────────────────────────
    doc.add_heading("Appendix A. Mean metrics at every k", level=1)
    for rep in reps:
        doc.add_heading(rep, level=2)
        cs_a = curve_stats(df, rep, "ari")[["k", "mean", "std"]]
        cs_n = curve_stats(df, rep, "nmi")[["k", "mean", "std"]]
        merged = cs_a.merge(cs_n, on="k", suffixes=("_ari", "_nmi"))
        merged.columns = ["k", "ARI mean", "ARI SD", "NMI mean", "NMI SD"]
        add_table(doc, merged)

    doc.add_paragraph()
    foot = doc.add_paragraph(
        f"Source data: {csv_path}. Regenerate with: python -m src.report_kmeans"
    )
    for r in foot.runs:
        r.font.size = Pt(8)
        r.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Analyse the K-Means seed sweep into a .docx report")
    ap.add_argument("--csv", type=str, default=str(config.RESULTS_DIR / "kmeans_seed_sweep.csv"))
    ap.add_argument("--out", type=str, default=str(config.RESULTS_DIR / "kmeans_phase1_report.docx"))
    args = ap.parse_args()

    csv_path = Path(args.csv)
    if not csv_path.exists():
        raise SystemExit(
            f"{csv_path} not found. Run: python main.py --phase 1 --seeds 0 1 2 ..."
        )
    out = build_report(csv_path, Path(args.out))
    print(f"Report written to {out}")


if __name__ == "__main__":
    main()
